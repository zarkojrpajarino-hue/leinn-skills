#!/usr/bin/env python3
"""
Extract text from a document file for book-to-skill processing.

PDF extraction tries methods in order:
  1. pdftotext (poppler-utils) — best quality
  2. PyPDF2 — common Python library
  3. pdfminer.six — thorough fallback

EPUB extraction tries methods in order:
  1. ebooklib + BeautifulSoup4 — best quality
  2. zipfile + html.parser — stdlib fallback (no extra deps)

Other supported formats:
  - Plain text / Markdown / reStructuredText / AsciiDoc — direct read
  - HTML — BeautifulSoup4 if available, then stdlib html.parser
  - DOCX — python-docx if available, then stdlib ZIP/XML fallback
  - RTF — striprtf if available, then regex fallback
  - MOBI/AZW/AZW3 — Calibre ebook-convert if available

Outputs:
  <tempdir>/book_skill_work/full_text.txt  — full extracted text
  <tempdir>/book_skill_work/metadata.json  — stats and metadata

Set BOOK_SKILL_WORKDIR to override the output directory.
"""

import html
import html.parser
import importlib.util
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

OUTPUT_DIR = Path(
    os.environ.get(
        "BOOK_SKILL_WORKDIR",
        str(Path(tempfile.gettempdir()) / "book_skill_work"),
    )
)
OUTPUT_TEXT = OUTPUT_DIR / "full_text.txt"
OUTPUT_META = OUTPUT_DIR / "metadata.json"

WORDS_PER_TOKEN = 0.75  # approximate

TEXT_EXTENSIONS = {".txt", ".text", ".md", ".markdown", ".rst", ".adoc", ".asciidoc"}
HTML_EXTENSIONS = {".html", ".htm", ".xhtml"}
CALIBRE_EBOOK_EXTENSIONS = {".mobi", ".azw", ".azw3"}
SUPPORTED_EXTENSIONS = {
    ".pdf", ".epub", ".docx", ".rtf",
    *TEXT_EXTENSIONS,
    *HTML_EXTENSIONS,
    *CALIBRE_EBOOK_EXTENSIONS,
}

PYTHON_DEPENDENCIES = {
    "docling": "docling",
    "PyPDF2": "PyPDF2",
    "pdfminer": "pdfminer.six",
    "ebooklib": "ebooklib",
    "bs4": "beautifulsoup4",
    "docx": "python-docx",
    "striprtf": "striprtf",
}


def estimate_tokens(text: str) -> int:
    return int(len(text.split()) / WORDS_PER_TOKEN)


def supported_formats_message() -> str:
    return ", ".join(sorted(SUPPORTED_EXTENSIONS))


def python_module_available(module_name: str) -> bool:
    return importlib.util.find_spec(module_name) is not None


def missing_python_packages(module_names: list[str]) -> list[str]:
    missing = []
    for module_name in module_names:
        if not python_module_available(module_name):
            missing.append(PYTHON_DEPENDENCIES[module_name])
    return missing


def install_python_packages(packages: list[str]) -> bool:
    if not packages:
        return True

    print(f"Installing missing Python package(s): {', '.join(packages)}")
    try:
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", *packages],
            text=True,
            timeout=600,
        )
    except Exception as exc:
        print(f"Package installation failed: {exc}", file=sys.stderr)
        return False

    importlib.invalidate_caches()
    return result.returncode == 0


def normalize_install_mode(argv: list[str]) -> str:
    mode = os.environ.get("BOOK_SKILL_INSTALL_MISSING", "ask").lower()
    if "--no-install-missing" in argv:
        return "no"
    if "--install-missing" in argv:
        idx = argv.index("--install-missing")
        if idx + 1 < len(argv) and not argv[idx + 1].startswith("--"):
            mode = argv[idx + 1].lower()
        else:
            mode = "yes"
    if mode in {"1", "true", "y", "yes", "install"}:
        return "yes"
    if mode in {"0", "false", "n", "no", "fallback", "skip"}:
        return "no"
    return "ask"


def offer_dependency_install(
    *,
    feature: str,
    module_names: list[str],
    fallback: str | None,
    install_mode: str,
) -> None:
    packages = missing_python_packages(module_names)
    if not packages:
        return

    message = f"{feature} uses {', '.join(packages)} if installed"
    if fallback:
        message += f", otherwise {fallback}"
    message += "."
    print(message)

    should_install = False
    if install_mode == "yes":
        should_install = True
    elif install_mode == "ask" and sys.stdin.isatty():
        answer = input("Missing package(s) detected. Do you want to install? y=install, n=fallback: ").strip().lower()
        should_install = answer in {"y", "yes", "install"}
    else:
        if fallback:
            print("Non-interactive mode or install disabled; using fallback.")
        else:
            print("Non-interactive mode or install disabled; installation skipped.")

    if not should_install:
        if fallback:
            print(f"Using fallback: {fallback}.")
        return

    if install_python_packages(packages):
        still_missing = missing_python_packages(module_names)
        if not still_missing:
            print("Package installation complete.")
            return
        print(f"Package installation incomplete; still missing: {', '.join(still_missing)}", file=sys.stderr)
    else:
        print("Package installation failed.", file=sys.stderr)

    if fallback:
        print(f"Using fallback: {fallback}.")


def prepare_dependencies(ext: str, extraction_mode: str, install_mode: str) -> None:
    if ext == ".pdf" and extraction_mode == "technical":
        offer_dependency_install(
            feature="Technical PDF extraction",
            module_names=["docling"],
            fallback="the PDF text fallback chain",
            install_mode=install_mode,
        )

    if ext == ".pdf" and not shutil.which("pdftotext"):
        offer_dependency_install(
            feature="PDF text extraction",
            module_names=["PyPDF2", "pdfminer"],
            fallback="any installed Python PDF parser; extraction fails if none are available",
            install_mode=install_mode,
        )

    if ext == ".epub":
        offer_dependency_install(
            feature="EPUB extraction",
            module_names=["ebooklib", "bs4"],
            fallback="a stdlib ZIP/HTML parser",
            install_mode=install_mode,
        )

    if ext in HTML_EXTENSIONS:
        offer_dependency_install(
            feature="HTML extraction",
            module_names=["bs4"],
            fallback="a stdlib HTML parser",
            install_mode=install_mode,
        )

    if ext == ".docx":
        offer_dependency_install(
            feature="DOCX extraction",
            module_names=["docx"],
            fallback="a stdlib ZIP/XML parser",
            install_mode=install_mode,
        )

    if ext == ".rtf":
        offer_dependency_install(
            feature="RTF extraction",
            module_names=["striprtf"],
            fallback="a basic regex cleanup fallback",
            install_mode=install_mode,
        )


def read_text_file(path: str) -> str | None:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return Path(path).read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except Exception:
            return None
    return None


def extract_html_content(raw_html: str) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw_html, "html.parser")
        for element in soup(["script", "style", "head"]):
            element.decompose()
        return soup.get_text(separator="\n")
    except ImportError:
        parser = _HTMLTextExtractor()
        parser.feed(raw_html)
        return parser.get_text()


def extract_html_file(path: str) -> str | None:
    raw = read_text_file(path)
    if raw is None:
        return None
    return extract_html_content(raw)


def extract_docx_with_python_docx(docx_path: str) -> str | None:
    try:
        import docx
        document = docx.Document(docx_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except ImportError:
        return None
    except Exception:
        return None


def extract_docx_with_zipfile(docx_path: str) -> str | None:
    try:
        import xml.etree.ElementTree as ET

        with zipfile.ZipFile(docx_path) as zf:
            xml_bytes = zf.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        parts: list[str] = []
        for paragraph in root.iter(f"{namespace}p"):
            texts = [node.text for node in paragraph.iter(f"{namespace}t") if node.text]
            if texts:
                parts.append("".join(texts))
        return "\n".join(parts) if parts else None
    except Exception:
        return None


def extract_docx(docx_path: str) -> tuple[str, str]:
    print("Trying python-docx...", end=" ", flush=True)
    text = extract_docx_with_python_docx(docx_path)
    if text and text.strip():
        print("OK")
        return text, "python-docx"

    print("not available")
    print("Trying stdlib DOCX parser...", end=" ", flush=True)
    text = extract_docx_with_zipfile(docx_path)
    if text and text.strip():
        print("OK")
        return text, "zipfile-docx"

    print("FAILED")
    print(
        "\nERROR: Could not extract text from DOCX.\n"
        "Install python-docx for best results:\n"
        "  pip3 install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


def strip_rtf_fallback(raw: str) -> str:
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\par[d]?", "\n", raw)
    raw = re.sub(r"\\tab", "\t", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw)
    raw = raw.replace("{", "").replace("}", "")
    return html.unescape(raw)


def extract_rtf(rtf_path: str) -> tuple[str, str]:
    raw = read_text_file(rtf_path)
    if raw is None:
        print("ERROR: Could not read RTF file", file=sys.stderr)
        sys.exit(1)

    try:
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(raw)
        if text.strip():
            return text, "striprtf"
    except ImportError:
        pass
    except Exception:
        pass

    return strip_rtf_fallback(raw), "rtf-regex"


def extract_with_ebook_convert(input_path: str) -> str | None:
    if not shutil.which("ebook-convert"):
        return None
    output_path = OUTPUT_DIR / "ebook-convert-output.txt"
    try:
        result = subprocess.run(
            ["ebook-convert", input_path, str(output_path)],
            capture_output=True, text=True, timeout=300
        )
        if result.returncode == 0 and output_path.exists():
            text = output_path.read_text(encoding="utf-8", errors="replace")
            if text.strip():
                return text
    except Exception:
        pass
    return None


def extract_with_pdftotext(pdf_path: str) -> str | None:
    if not shutil.which("pdftotext"):
        return None
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", pdf_path, "-"],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except Exception:
        pass
    return None


def extract_with_pypdf2(pdf_path: str) -> str | None:
    try:
        import PyPDF2
        text_parts = []
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception:
                    text_parts.append("")
        return "\n".join(text_parts)
    except ImportError:
        return None
    except Exception:
        return None


def extract_with_pdfminer(pdf_path: str) -> str | None:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(pdf_path)
    except ImportError:
        return None
    except Exception:
        return None


def extract_with_ebooklib(epub_path: str) -> str | None:
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(epub_path)
        parts = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            parts.append(soup.get_text(separator="\n"))
        return "\n\n".join(parts)
    except ImportError:
        return None
    except Exception:
        return None


class _HTMLTextExtractor(html.parser.HTMLParser):
    """Minimal HTML → plain text converter using stdlib only."""

    SKIP_TAGS = {"script", "style", "head"}

    def __init__(self):
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth = 0
        self._current_skip: str | None = None

    def handle_starttag(self, tag, attrs):
        if tag in self.SKIP_TAGS:
            self._skip_depth += 1
        if tag in ("p", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "div"):
            self._parts.append("\n")

    def handle_endtag(self, tag):
        if tag in self.SKIP_TAGS and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data):
        if not self._skip_depth:
            self._parts.append(data)

    def get_text(self) -> str:
        return html.unescape("".join(self._parts))


def extract_with_zipfile(epub_path: str) -> str | None:
    """stdlib-only EPUB extractor: unzip → parse HTML files."""
    try:
        with zipfile.ZipFile(epub_path) as zf:
            names = zf.namelist()
            # Read OPF spine to get reading order, fall back to sorted xhtml files
            spine_order: list[str] = []
            opf_files = [n for n in names if n.endswith(".opf")]
            if opf_files:
                opf_text = zf.read(opf_files[0]).decode("utf-8", errors="replace")
                spine_order = re.findall(r'href=["\']([^"\']+\.(?:xhtml|html))["\']', opf_text)

            html_files = spine_order or sorted(
                n for n in names if n.endswith((".html", ".xhtml"))
            )
            if not html_files:
                return None

            parts = []
            for name in html_files:
                try:
                    raw = zf.read(name).decode("utf-8", errors="replace")
                    parser = _HTMLTextExtractor()
                    parser.feed(raw)
                    parts.append(parser.get_text())
                except Exception:
                    continue
            return "\n\n".join(parts) if parts else None
    except Exception:
        return None


def extract_epub(epub_path: str) -> tuple[str, str]:
    """Return (text, method) for an EPUB file."""
    print("Trying ebooklib + BeautifulSoup4...", end=" ", flush=True)
    text = extract_with_ebooklib(epub_path)
    if text and text.strip():
        print("OK")
        return text, "ebooklib"

    print("not available")
    print("Trying stdlib zipfile parser...", end=" ", flush=True)
    text = extract_with_zipfile(epub_path)
    if text and text.strip():
        print("OK")
        return text, "zipfile"

    print("FAILED")
    print(
        "\nERROR: Could not extract text from EPUB.\n"
        "Install ebooklib + beautifulsoup4 for best results:\n"
        "  pip3 install ebooklib beautifulsoup4",
        file=sys.stderr,
    )
    sys.exit(1)


def count_epub_chapters(epub_path: str) -> int:
    """Count spine items (approximate chapter count) without dependencies."""
    try:
        with zipfile.ZipFile(epub_path) as zf:
            opf_files = [n for n in zf.namelist() if n.endswith(".opf")]
            if not opf_files:
                return 0
            opf_text = zf.read(opf_files[0]).decode("utf-8", errors="replace")
            return len(re.findall(r'<itemref\b', opf_text))
    except Exception:
        return 0


def count_pages(pdf_path: str) -> int:
    # Try pdfinfo first
    if shutil.which("pdfinfo"):
        try:
            result = subprocess.run(
                ["pdfinfo", pdf_path], capture_output=True, text=True, timeout=15
            )
            for line in result.stdout.splitlines():
                if line.startswith("Pages:"):
                    return int(line.split(":")[1].strip())
        except Exception:
            pass
    # Fallback: count form-feed chars (pdftotext -layout uses \f between pages)
    try:
        import PyPDF2
        with open(pdf_path, "rb") as f:
            return len(PyPDF2.PdfReader(f).pages)
    except Exception:
        return 0


def detect_structure(text: str) -> dict:
    """Detect chapter count and table of contents presence."""
    import re
    lines = text[:50000].splitlines()

    # Look for chapter headings
    chapter_pattern = re.compile(
        r"^\s*(chapter\s+\d+|CHAPTER\s+\d+|ch\.\s*\d+|\d+\.\s+[A-Z])",
        re.IGNORECASE
    )
    chapters_found = [l.strip() for l in lines if chapter_pattern.match(l)]

    # Look for ToC indicators in the first ~30k chars (front matter is often well past 5k:
    # copyright pages, praise, dedications, forewords). Require the keyword to appear on
    # its own line to avoid false positives like "the contents of this book are...".
    toc_pattern = re.compile(
        r"^\s*(?:table of contents|contents|índice|sumário)\s*$",
        re.IGNORECASE | re.MULTILINE,
    )
    has_toc = bool(toc_pattern.search(text[:30000]))

    return {
        "chapters_detected": len(chapters_found),
        "chapter_headings_sample": chapters_found[:10],
        "has_toc": has_toc,
    }


def extract_with_docling(pdf_path: str) -> str | None:
    """Layout-aware extraction using Docling. Best for technical books with tables and code."""
    try:
        from docling.document_converter import DocumentConverter
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.datamodel.base_models import InputFormat
        from docling.document_converter import PdfFormatOption

        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = False
        pipeline_options.do_table_structure = True

        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        result = converter.convert(pdf_path)
        return result.document.export_to_markdown()
    except ImportError:
        return None
    except Exception:
        return None


def main():
    if len(sys.argv) < 2:
        print("Usage: extract.py <path-to-document> [--mode technical|text] [--install-missing ask|yes|no]", file=sys.stderr)
        print(f"Supported formats: {supported_formats_message()}", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]
    install_mode = normalize_install_mode(sys.argv)

    # Parse --mode flag
    extraction_mode = "text"
    if "--mode" in sys.argv:
        idx = sys.argv.index("--mode")
        if idx + 1 < len(sys.argv):
            extraction_mode = sys.argv[idx + 1].lower()
    if extraction_mode not in ("technical", "text"):
        extraction_mode = "text"

    if not os.path.exists(input_path):
        print(f"ERROR: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    input_file = Path(input_path)
    ext = input_file.suffix.lower()
    document_format = ext.lstrip(".")

    # Sniff magic bytes as a fallback for files without useful extensions.
    if ext not in SUPPORTED_EXTENSIONS:
        with open(input_path, "rb") as f:
            header = f.read(8)
        if header[:4] == b"%PDF":
            ext = ".pdf"
            document_format = "pdf"
        elif header[:2] == b"PK":
            try:
                with zipfile.ZipFile(input_path) as zf:
                    names = set(zf.namelist())
                    if "mimetype" in names and zf.read("mimetype").startswith(b"application/epub"):
                        ext = ".epub"
                        document_format = "epub"
                    elif "word/document.xml" in names:
                        ext = ".docx"
                        document_format = "docx"
                    else:
                        print(
                            f"ERROR: Unsupported ZIP-based format '{input_file.name}'. Supported: {supported_formats_message()}",
                            file=sys.stderr,
                        )
                        sys.exit(1)
            except (zipfile.BadZipFile, KeyError, OSError):
                print(
                    f"ERROR: Unsupported ZIP-based format '{input_file.name}'. Supported: {supported_formats_message()}",
                    file=sys.stderr,
                )
                sys.exit(1)
        else:
            print(
                f"ERROR: Unsupported format '{ext or '<none>'}'. Supported: {supported_formats_message()}",
                file=sys.stderr,
            )
            sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    prepare_dependencies(ext, extraction_mode, install_mode)

    if ext in CALIBRE_EBOOK_EXTENSIONS and not shutil.which("ebook-convert"):
        print(
            "MOBI/AZW/AZW3 extraction requires Calibre's ebook-convert command. "
            "Install Calibre and ensure ebook-convert is on PATH, then rerun this command.",
            file=sys.stderr,
        )
        sys.exit(1)

    if ext == ".epub":
        print(f"Extracting EPUB: {input_path}")
        text, method = extract_epub(input_path)
        pages = count_epub_chapters(input_path)
        pages_label = "spine_items"
    elif ext == ".pdf":
        print(f"Extracting PDF: {input_path}")
        if extraction_mode == "technical":
            print("Mode: technical — using Docling (layout-aware)...", end=" ", flush=True)
            text = extract_with_docling(input_path)
            if text:
                method = "docling"
                print("OK")
            else:
                print("not available, falling back to pdftotext")
                extraction_mode = "text"

        if extraction_mode == "text":
            print("Mode: text — using pdftotext...")
            print("Trying pdftotext...", end=" ", flush=True)
            text = extract_with_pdftotext(input_path)

            if text:
                method = "pdftotext"
                print("OK")
            else:
                print("not available")
                print("Trying PyPDF2...", end=" ", flush=True)
                text = extract_with_pypdf2(input_path)
                if text:
                    method = "PyPDF2"
                    print("OK")
                else:
                    print("not available")
                    print("Trying pdfminer.six...", end=" ", flush=True)
                    text = extract_with_pdfminer(input_path)
                    if text:
                        method = "pdfminer"
                        print("OK")
                    else:
                        print("FAILED")
                        print(
                            "\nERROR: Could not extract text from PDF.\n"
                            "Install one of: poppler-utils (pdftotext), PyPDF2, or pdfminer.six\n"
                            "  sudo apt install poppler-utils\n"
                            "  pip3 install PyPDF2\n"
                            "  pip3 install pdfminer.six",
                            file=sys.stderr,
                        )
                        sys.exit(1)

        pages = count_pages(input_path)
        pages_label = "pages"
    elif ext in TEXT_EXTENSIONS:
        print(f"Extracting text document: {input_path}")
        text = read_text_file(input_path)
        if text is None or not text.strip():
            print("ERROR: Could not read text document", file=sys.stderr)
            sys.exit(1)
        method = "plain-text"
        pages = 0
        pages_label = "sections"
    elif ext in HTML_EXTENSIONS:
        print(f"Extracting HTML: {input_path}")
        text = extract_html_file(input_path)
        if text is None or not text.strip():
            print("ERROR: Could not extract text from HTML", file=sys.stderr)
            sys.exit(1)
        method = "html-parser"
        pages = 0
        pages_label = "sections"
    elif ext == ".docx":
        print(f"Extracting DOCX: {input_path}")
        text, method = extract_docx(input_path)
        pages = 0
        pages_label = "sections"
    elif ext == ".rtf":
        print(f"Extracting RTF: {input_path}")
        text, method = extract_rtf(input_path)
        pages = 0
        pages_label = "sections"
    elif ext in CALIBRE_EBOOK_EXTENSIONS:
        print(f"Extracting ebook with Calibre: {input_path}")
        text = extract_with_ebook_convert(input_path)
        if text is None or not text.strip():
            print(
                f"ERROR: Could not extract text from {ext}. Install Calibre and ensure ebook-convert is on PATH.",
                file=sys.stderr,
            )
            sys.exit(1)
        method = "ebook-convert"
        pages = 0
        pages_label = "sections"
    else:
        print(
            f"ERROR: Unsupported format '{ext}'. Supported: {supported_formats_message()}",
            file=sys.stderr,
        )
        sys.exit(1)

    # Write full text
    OUTPUT_TEXT.write_text(text, encoding="utf-8")

    tokens = estimate_tokens(text)
    structure = detect_structure(text)
    file_size_mb = os.path.getsize(input_path) / (1024 * 1024)

    metadata = {
        "source_file": str(Path(input_path).resolve()),
        "filename": Path(input_path).name,
        "format": document_format,
        "extraction_method": method,
        "extraction_mode": extraction_mode,
        "file_size_mb": round(file_size_mb, 2),
        pages_label: pages,
        "chars": len(text),
        "words": len(text.split()),
        "estimated_tokens": tokens,
        "estimated_tokens_human": f"~{tokens // 1000}K",
        "output_text": str(OUTPUT_TEXT),
        **structure,
    }

    OUTPUT_META.write_text(json.dumps(metadata, indent=2, ensure_ascii=False))

    page_label = {
        "spine_items": "Spine items",
        "pages": "Pages",
        "sections": "Sections",
    }.get(pages_label, pages_label.replace("_", " ").title())
    page_line = f"   {page_label}: {pages}"
    print("\nExtraction complete:")
    print(f"   Format  : {document_format.upper()}")
    print(f"   Method  : {method}")
    print(page_line)
    print(f"   Words   : {len(text.split()):,}")
    print(f"   Tokens  : ~{tokens // 1000}K")
    print(f"   Chapters: {structure['chapters_detected']} detected")
    print(f"   ToC     : {'yes' if structure['has_toc'] else 'not detected'}")
    if not structure["has_toc"]:
        print(
            "   WARN    : No table of contents detected — chapter mapping in Step 3 "
            "will rely on heading scan only, which may miss or duplicate sections."
        )
    print(f"\n   Text -> {OUTPUT_TEXT}")
    print(f"   Meta -> {OUTPUT_META}")


if __name__ == "__main__":
    main()
